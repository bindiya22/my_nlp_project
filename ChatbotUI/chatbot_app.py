import streamlit as st
import PyPDF2
import docx
import os

def send_message(loaded_document):
    user_message = st.session_state.user_input
    st.session_state.chat_history.append(("You", user_message))

    chatbot_response = process_document(loaded_document) if loaded_document else "Please load a document first."

    st.session_state.chat_history.append(("ChatbotName", chatbot_response))
    st.session_state.user_input = ""

    for sender, message in st.session_state.chat_history:
        with st.chat_message(sender):
            st.write(message)

def process_document(document_content):
    # ***IMPORTANT: Replace this with your actual document processing logic***
    # This is a placeholder example:
    if "important" in document_content.lower():
        return "Document mentions 'important'."
    elif "summary" in document_content.lower():
        return "Document mentions 'summary'"
    else:
        return "Document processed (example response). Implement your analysis here!"

def load_text(file_path):
    try:
        with open(file_path, 'r', encoding="utf-8") as file:
            extracted_text = file.read()
            st.write(f"Loaded Text:\n{extracted_text}\n")
            global loaded_document
            loaded_document = extracted_text
    except Exception as e:
        st.error(f"Error loading text: {e}\n")

def load_document(file_path):
    document_content = ""
    try:
        file_path = os.path.abspath(file_path)

        if file_path.endswith('.txt'):
            try:
                with open(file_path, 'r', encoding="utf-8") as file:
                    document_content = file.read()
            except UnicodeDecodeError:
                with open(file_path, 'r', encoding="latin-1") as file:
                    document_content = file.read()
            except Exception as e:
                st.error(f"Error reading TXT: {e}")

        elif file_path.endswith('.pdf'):
            try:
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    all_text = []
                    for page in reader.pages:
                        text = page.extract_text()
                        if text:
                            all_text.append(text)
                    document_content = "\n".join(all_text)
            except Exception as e:
                st.error(f"Error reading PDF: {e}")

        elif file_path.endswith('.docx'):
            try:
                doc = docx.Document(file_path)
                all_text = []
                for paragraph in doc.paragraphs:
                    if paragraph.text:
                        all_text.append(paragraph.text)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text:
                                all_text.append(cell.text)
                document_content = "\n".join(all_text)
            except Exception as e:
                st.error(f"Error reading DOCX: {e}")

        else:
            document_content = "Unsupported file type."

    except Exception as e:
        st.error(f"Error loading document: {e}")

    st.write(f"Loaded Document:\n{document_content}\n")

    global loaded_document
    loaded_document = document_content


st.title("Docubot")

if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "user_input" not in st.session_state:
    st.session_state.user_input = ""

uploaded_file = st.file_uploader("Choose a file", type=["txt", "pdf", "docx"])

if uploaded_file is not None:
    file_path = uploaded_file.name
    with open(file_path, 'wb') as f:
        f.write(uploaded_file.getbuffer())
    if file_path.endswith('.txt'):
        load_text(file_path)
    else:
        load_document(file_path)

user_input = st.text_input("Enter your message:", key="user_input")

if st.button("Send"):
    if 'loaded_document' in globals() and loaded_document:
        send_message(loaded_document)
    else:
        st.write("Please load a document first.")