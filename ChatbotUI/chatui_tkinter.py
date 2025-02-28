import tkinter as tk
from tkinter import scrolledtext, filedialog
import PyPDF2
import docx
import os

def send_message():
    user_message = user_input.get()
    chat_display.insert(tk.END, "You: " + user_message + "\n")
    user_input.delete(0, tk.END)

    if 'loaded_document' in globals() and loaded_document:
        chatbot_response = process_document(loaded_document)
    else:
        chatbot_response = "Please load a document first."  # More informative message
    chat_display.insert(tk.END, "ChatbotName: " + chatbot_response + "\n")

def process_document(document_content):
    # ***IMPORTANT: Replace this with your actual document processing logic***
    # This is a placeholder example:
    if "important" in document_content.lower():
        return "Document mentions 'important'."
    elif "summary" in document_content.lower():
        return "Document mentions 'summary'"
    else:
        return "Document processed (example response).  Implement your analysis here!"

def load_text():
    file_path = filedialog.askopenfilename(filetypes=[("Text files", "*.txt")])
    if file_path:
        try:
            with open(file_path, 'r', encoding="utf-8") as file:
                extracted_text = file.read()
                chat_display.insert(tk.END, f"Loaded Text:\n{extracted_text}\n")
                global loaded_document
                loaded_document = extracted_text
        except Exception as e:
            chat_display.insert(tk.END, f"Error loading text: {e}\n")

def load_document():
    file_path = filedialog.askopenfilename(filetypes=[
        ("Text files", "*.txt"),
        ("PDF files", "*.pdf"),
        ("Word files", "*.docx"),
    ])

    if not file_path:
        return

    document_content = ""

    try:
        file_path = os.path.abspath(file_path)  # Convert to absolute path

        if file_path.endswith('.txt'):
            try:
                with open(file_path, 'r', encoding="utf-8") as file:
                    document_content = file.read()
            except UnicodeDecodeError: # Handle potential encoding errors
                with open(file_path, 'r', encoding="latin-1") as file: # Try a different encoding
                    document_content = file.read()
            except Exception as e:
                document_content = f"Error reading TXT: {e}"

        elif file_path.endswith('.pdf'):
            try:
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    all_text = [] # More efficient way to collect text
                    for page in reader.pages:
                        text = page.extract_text()
                        if text:  # The KEY FIX for PDFs!
                            all_text.append(text)
                    document_content = "\n".join(all_text) # Join all text at once
            except Exception as e:
                document_content = f"Error reading PDF: {e}"

        elif file_path.endswith('.docx'):
            try:
                doc = docx.Document(file_path)
                all_text = [] # More efficient way to collect text
                for paragraph in doc.paragraphs:
                    if paragraph.text:  # Check for None or ""
                        all_text.append(paragraph.text)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text:  # Check for None or ""
                                all_text.append(cell.text)
                document_content = "\n".join(all_text) # Join all text at once
            except Exception as e:
                document_content = f"Error reading DOCX: {e}"

        else:
            document_content = "Unsupported file type."

    except Exception as e:
        document_content = f"Error loading document: {e}"

    chat_display.insert(tk.END, f"Loaded Document:\n{document_content}\n")

    global loaded_document
    loaded_document = document_content
root = tk.Tk()
root.title("Docubot")

# Use grid layout
root.grid_rowconfigure(0, weight=1)  # Make row 0 (chat display) expandable
root.grid_columnconfigure(0, weight=1) # Make column 0 (chat display) expandable

chatbot_name_label = tk.Label(root, text="Welcome to DOCUBOT", font=("Helvetica", 16))
chatbot_name_label.grid(row=0, column=0, sticky="ew", pady=10) # sticky makes it fill width

chat_display = scrolledtext.ScrolledText(root, wrap=tk.WORD)  # Removed fixed width/height
chat_display.grid(row=1, column=0, sticky="nsew", padx=10, pady=10) # sticky makes it fill space

input_frame = tk.Frame(root)
input_frame.grid(row=2, column=0, sticky="ew", padx=10, pady=5) # sticky makes it fill width

user_input = tk.Entry(input_frame)
user_input.grid(row=0, column=0, sticky="ew", padx=5) # sticky makes it fill space
input_frame.grid_columnconfigure(0, weight=1) # Make entry expand

send_button = tk.Button(input_frame, text="Send", command=send_message)
send_button.grid(row=0, column=1, padx=5)

button_frame = tk.Frame(root)
button_frame.grid(row=3, column=0, sticky="ew", padx=10, pady=5) # sticky makes it fill width

load_text_button = tk.Button(button_frame, text="Load Text", command=load_text)
load_text_button.grid(row=0, column=0, padx=5)

load_doc_button = tk.Button(button_frame, text="Load Document", command=load_document)
load_doc_button.grid(row=0, column=1, padx=5)

root.mainloop()